from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path
import tempfile

from Agent.app.config import AppConfig
from Agent.app.llm.openai_client import chat_complete, create_openai_client


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--contract-type", default="技术服务合同")
    parser.add_argument("--company-a", default="甲方有限公司")
    parser.add_argument("--company-b", default="乙方科技有限公司")
    parser.add_argument("--sign-date", default=date.today().isoformat())
    parser.add_argument("--model")
    parser.add_argument("--output", required=True)
    parser.add_argument("--temperature", type=float, default=0.3)
    args = parser.parse_args()

    output_path = Path(args.output).resolve()
    if output_path.suffix.lower() not in {".docx", ".pdf"}:
        raise ValueError("Output must be .docx or .pdf")

    system_prompt = _build_system_prompt()
    user_prompt = _build_user_prompt(
        contract_type=args.contract_type,
        company_a=args.company_a,
        company_b=args.company_b,
        sign_date=args.sign_date,
    )
    config = AppConfig.from_env(legal_workspace=Path.cwd(), model=args.model)
    client = create_openai_client(
        api_key=config.openai_api_key,
        base_url=config.openai_base_url,
    )
    contract_text = chat_complete(
        client=client,
        model=config.model,
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        temperature=args.temperature,
    )
    if output_path.suffix.lower() == ".docx":
        _save_docx(contract_text, output_path)
        return
    _save_pdf(contract_text, output_path)


def _build_system_prompt() -> str:
    return (
        "你是合同起草助手，负责生成中文商业合同。"
        "合同要正式、清晰、可直接用于业务沟通。"
        "输出纯文本，不要使用Markdown或列表符号。"
    )


def _build_user_prompt(
    contract_type: str,
    company_a: str,
    company_b: str,
    sign_date: str,
) -> str:
    return (
        "请生成一份中文商业合同，满足以下要求：\n"
        f"1) 合同类型：{contract_type}\n"
        f"2) 甲方：{company_a}\n"
        f"3) 乙方：{company_b}\n"
        f"4) 签署日期：{sign_date}\n"
        "5) 包含合同编号、合同期限、服务/供货范围、价格与支付、"
        "交付与验收、保密、知识产权、违约责任、争议解决、其他条款\n"
        "6) 条款要具体、可执行，字数不少于1200字\n"
        "7) 输出必须是带换行的正文，分段清晰，标题在第一行\n"
    )


def _save_docx(content: str, output_path: Path) -> None:
    from docx import Document

    document = Document()
    for line in content.splitlines():
        if line.strip():
            document.add_paragraph(line.strip())
        else:
            document.add_paragraph("")
    document.save(str(output_path))


def _save_pdf(content: str, output_path: Path) -> None:
    try:
        from docx2pdf import convert
    except Exception as exc:
        raise RuntimeError("PDF生成需要安装docx2pdf依赖") from exc
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_docx = Path(temp_dir) / "contract.docx"
        _save_docx(content, temp_docx)
        convert(str(temp_docx), str(output_path))


if __name__ == "__main__":
    main()
